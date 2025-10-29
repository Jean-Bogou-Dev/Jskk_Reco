from flask import Flask, render_template, request, redirect, url_for, session
import mysql.connector
import base64
import cv2
import numpy as np
import datetime
import os
from datetime import datetime, date
from docx import Document
from flask import send_file
import io

app = Flask(__name__)
app.secret_key = 'votre_cle_secrete'

def get_db_connection():
    return mysql.connector.connect(
        host="localhost",
        user="root",
        password="",
        database="reconnaisancefacial"
    )

# Page d'authentification
@app.route('/')
def login_page():
    return render_template('authentification.html')

login_attempts = {}

#formulaire de connexion
@app.route('/login', methods=['POST'])
def login():
    username = request.form['username']
    password = request.form['password']

    conn = get_db_connection()
    cursor = conn.cursor()
    #verification si l'utilisateur exist
    cursor.execute("SELECT * FROM USERS WHERE USERNAME=%s", (username,))
    user_exists = cursor.fetchone()

#vider le résultat précédent avant nouvelle requête
    cursor.fetchall()

    cursor.execute("SELECT * FROM USERS WHERE USERNAME=%s AND PASSWORD=%s", (username, password))
    user = cursor.fetchone()
    cursor.close()
    conn.close()

#si l'utilisateur est reconnu on redirectionne vers la page d'accueil "index.html"
    if user:
        session['username'] = username
        login_attempts[username] = 0
        return redirect(url_for('accueil'))
    
#si l'utilisateur se trompe de mot de passe plus de 3 fois le lien vers la page de recupération de mot de passe apparait
    else:
        login_attempts[username] = login_attempts.get(username, 0) + 1
        show_reset = False
        if login_attempts[username] >= 3 and user_exists:
            show_reset = True
        return render_template('authentification.html', error=True, show_reset=show_reset)

#Création d'un nouveau compte
@app.route('/create_account', methods=['POST'])
def create_account():
    username = request.form['new_username']
    password = request.form['new_password']
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("INSERT INTO USERS (USERNAME, PASSWORD) VALUES (%s, %s)", (username, password))
    conn.commit()
    cursor.close()
    conn.close()
    return redirect(url_for('login_page'))

#récupération du mot de passe grace au nom d'utilisateur
@app.route('/reset_password', methods=['POST'])
def reset_password():
    username = request.form['reset_username']

    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT PASSWORD FROM USERS WHERE USERNAME=%s", (username,))
    result = cursor.fetchone()
    cursor.fetchall()
    cursor.close()
    conn.close()

    if result:
        recovered_password = result[0]
        return render_template('authentification.html',
                               recovered_user=username,
                               recovered_password=recovered_password)
    else:
        return render_template('authentification.html',
                               show_create=True)



# Page d'accueil
@app.route('/accueil')
def accueil():
    if 'username' not in session:
        return redirect(url_for('login_page'))
    return render_template('index.html', username=session['username'])

# Page détail étudiant
@app.route('/etudiant')
def detail_etudiant():
    if 'username' not in session:
        return redirect(url_for('login_page'))
    return render_template('detail_etudiant.html')

#route d'ajout d'etudiant innexistant depuis présence
@app.route('/ajouter_etudiant')
def ajouter_etudiant():
    return render_template('detail_etudiant.html')


@app.route('/ajouter', methods=['POST'])

def ajouter():
       # Récupération des données du formulaire
    CNEMASSAR = request.form['CNEMASSAR']
    nom = request.form['nom']
    prenom = request.form['prenom']
    sex = request.form['sex']
    telephone = request.form['telephone']
    mail = request.form['mail']
    departement = request.form['departement']
    chef_departement = request.form['chef_departement']
    options = request.form['options']
    annee_academique = request.form['annee_academique']
    semestre = request.form['semestre']
    photo_data = request.form.get('photo_data')

    conn = get_db_connection()
    cursor = conn.cursor()

    # Insérer l'étudiant
    cursor.execute("""INSERT INTO etudiant (
                   CNE_MASSAR, 
                   nom, 
                   prenom, 
                   sex, 
                   telephone, 
                   mail, 
                   options, 
                   annee_academique,
                   semestre) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)""", 
                   (CNEMASSAR, 
                   nom, 
                   prenom, 
                   sex, 
                   telephone, 
                   mail, 
                   options, 
                   annee_academique,
                   semestre))
    cursor.execute("""INSERT INTO departement (
                   nom_departement, 
                   chef_departement,CNE_MASSAR) VALUES (%s, %s, %s)""", 
                   ( departement, 
                   chef_departement,CNEMASSAR))

    # Traitement de la photo
    if photo_data:
        header, encoded = photo_data.split(",", 1)
        image_bytes = base64.b64decode(encoded)
        nparr = np.frombuffer(image_bytes, np.uint8)
        img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

        # Détection du visage
        face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        faces = face_cascade.detectMultiScale(gray, 1.3, 5)

        if len(faces) > 0:
            x, y, w, h = faces[0]
            face_img = img[y:y+h, x:x+w]

            # Conversion en RGB pour compatibilité avec face_recognition
            rgb_face_img = cv2.cvtColor(face_img, cv2.COLOR_BGR2RGB)

            os.makedirs("static/photos", exist_ok=True)
            # Sauvegarde dans un fichier
            filename = f"static/photos/{CNEMASSAR}.jpg"
            success = cv2.imwrite(filename, rgb_face_img)
            print("Image enregistrée dans le fichier :", success)

                        # Encodage unique de l’image RGB
            _, buffer = cv2.imencode('.jpg', rgb_face_img)
            photo_data = buffer.tobytes()

            
            cursor.execute("SELECT CNE_MASSAR FROM IMAGE WHERE CNE_MASSAR = %s", (CNEMASSAR,))
            existing_image = cursor.fetchone()

            if existing_image:
                cursor.execute("UPDATE IMAGE SET IMAGE = %s WHERE CNE_MASSAR = %s", (photo_data, CNEMASSAR))
            else:
                cursor.execute("INSERT INTO IMAGE (CNE_MASSAR, IMAGE) VALUES (%s, %s)", (CNEMASSAR, photo_data))
    
    

    conn.commit()
    cursor.close()
    conn.close()

    return render_template('detail_etudiant.html',
                       success=True,
                       CNEMASSAR=CNEMASSAR,
                       nom=nom,
                       prenom=prenom,
                       sex=sex,
                       telephone=telephone,
                       mail=mail,
                       departement=departement,
                       chef_departement=chef_departement,
                       options=options,
                       annee_academique=annee_academique,
                       semestre=semestre)


# Page présence
from datetime import datetime, timedelta

@app.route('/presence', methods=['GET', 'POST'])
def presence():
    if 'username' not in session:
        return redirect(url_for('login_page'))

    infos = {}
    presences = []
    nb_presence = 0
    nb_absence = 0
    date_debut = ''
    date_fin = ''

    conn = get_db_connection()
    cursor = conn.cursor()

    # Affichage par défaut : tous les étudiants et leurs présences
    if request.method == 'GET':
        cursor.execute("""
            SELECT e.CNE_MASSAR, e.nom, e.prenom, d.nom_departement, e.mail, p.DATE_PRESENCE, p.HEURE, p.STATUT
            FROM etudiant e
            JOIN departement d ON e.CNE_MASSAR = d.CNE_MASSAR
            JOIN presence p ON e.CNE_MASSAR = p.CNE_MASSAR
            ORDER BY p.DATE_PRESENCE DESC
        """)
        presences = cursor.fetchall()

    # Traitement après soumission du formulaire
    elif request.method == 'POST':
        cne = request.form.get('cne')
        date_debut = request.form.get('date_debut')
        date_fin = request.form.get('date_fin')
        bouton = request.form.get('rapport') or request.form.get('Calculer')

        # Recherche d’un étudiant
        if cne:
            cursor.execute("SELECT nom, prenom, mail FROM etudiant WHERE CNE_MASSAR = %s", (cne,))
            result = cursor.fetchone()
            if result:
                infos = {
                    'cne': cne,
                    'nom': result[0],
                    'prenom': result[1],
                    'mail': result[2],
                    'statut': 'Enregistré' 
                }


                # Si bouton "Calculer" est cliqué on calcule les présences/absences
                action = request.form.get('action')

                #si la periode date debut et date fin sont choisi on calcule et on affiche dans le tableau les jours selectionner
                if action == 'Calculer' and cne and date_debut and date_fin:
                    #  Requête filtrée par période
                    cursor.execute("""
                        SELECT e.CNE_MASSAR, e.nom, e.prenom, d.nom_departement, e.mail, p.DATE_PRESENCE, p.HEURE, p.STATUT
                        FROM etudiant e
                        JOIN departement d ON e.CNE_MASSAR = d.CNE_MASSAR
                        JOIN presence p ON e.CNE_MASSAR = p.CNE_MASSAR
                        WHERE e.CNE_MASSAR = %s AND p.DATE_PRESENCE BETWEEN %s AND %s
                        ORDER BY p.DATE_PRESENCE DESC
                    """, (cne, date_debut, date_fin))
                    presences = cursor.fetchall()

                    #  Calcul des jours ouvrables
                    start = datetime.strptime(date_debut, "%Y-%m-%d")
                    end = datetime.strptime(date_fin, "%Y-%m-%d")
                    jours_ouvrables = sum(1 for i in range((end - start).days + 1)
                                        if (start + timedelta(days=i)).weekday() < 5)

                    nb_presence = sum(1 for p in presences if p[7] == 'Présent')
                    nb_absence = max(0, jours_ouvrables - nb_presence)

                    

                elif action == 'Rapport Etudiant(e)' and cne:
                    #  Requête complète sans filtre
                    cursor.execute("""
                        SELECT e.CNE_MASSAR, e.nom, e.prenom, d.nom_departement, e.mail, p.DATE_PRESENCE, p.HEURE, p.STATUT
                        FROM etudiant e
                        JOIN departement d ON e.CNE_MASSAR = d.CNE_MASSAR
                        JOIN presence p ON e.CNE_MASSAR = p.CNE_MASSAR
                        WHERE e.CNE_MASSAR = %s
                        ORDER BY p.DATE_PRESENCE DESC
                    """, (cne,))
                    presences = cursor.fetchall()



    cursor.close()
    conn.close()

    return render_template('presence.html',
                       infos=infos,
                       presences=presences,
                       nb_presence=nb_presence,
                       nb_absence=nb_absence,
                       date_debut=date_debut,
                       date_fin=date_fin)

@app.route('/rechercher_etudiant', methods=['POST'])
def rechercher_etudiant():
    if 'username' not in session:
        return redirect(url_for('login_page'))

    cne = request.form.get('cne_recherche')

    conn = get_db_connection()
    cursor = conn.cursor()

    # Vérifier si l'étudiant existe
    cursor.execute("SELECT nom, prenom, mail FROM etudiant WHERE CNE_MASSAR = %s", (cne,))
    result = cursor.fetchone()

    if result:
        infos = {
            'cne': cne,
            'nom': result[0],
            'prenom': result[1],
            'mail': result[2],
            'statut': 'Enregistré'
        }

        # Pas de période = pas de calcul de présence
        presences = []
        nb_presence = 0
        nb_absence = 0

        cursor.close()
        conn.close()

        return render_template('presence.html',
                               infos=infos,
                               presences=presences,
                               nb_presence=nb_presence,
                               nb_absence=nb_absence)
    else:
        cursor.close()
        conn.close()
        # Étudiant non trouvé = afficher popup
        return render_template('presence.html',
                       popup_cne=cne,
                       infos={},
                       presences=[],
                       nb_presence=0,
                       nb_absence=0,
                       date_debut='',
                       date_fin='')




# Page entraînement du modèle
@app.route('/entrainement')
def entrainement():
    if 'username' not in session:
        return redirect(url_for('login_page'))
    return render_template('entrainement.html')

@app.route('/lancer_entrainement', methods=['POST'])
def lancer_entrainement():
    try:
        import face_recognition

        encodings = []
        ids = []

        photos_path = "static/photos"
        for filename in os.listdir(photos_path):
            if filename.endswith(".jpg"):
                cne = filename.split(".")[0]
                image_path = os.path.join(photos_path, filename)

                # Chargement avec OpenCV pour contrôle du format
                img = cv2.imread(image_path)

                if img is None:
                    print(f"❌ Image introuvable : {filename}")
                    continue

                if len(img.shape) != 3 or img.shape[2] != 3:
                    print(f"❌ Format non pris en charge : {filename}")
                    continue

                # Conversion en RGB
                rgb_img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)

                # Détection et encodage
                print(f"📷 Traitement de {filename}")
                print("Shape:", rgb_img.shape)
                print("Dtype:", rgb_img.dtype)

                face_locations = face_recognition.face_locations(rgb_img)           #extraction des caracteristique du visage
                face_encs = face_recognition.face_encodings(rgb_img, face_locations)    #Encodage des caracteristiques dans une variable

                if face_encs:
                    encodings.append(face_encs[0])
                    ids.append(cne)
                    print(f"✅ Encodage réussi pour {cne}")
                else:
                    print(f"⚠️ Aucun visage détecté dans {filename}")

        # Sauvegarde des modèles dans des fichier
        np.save("model_encodings.npy", encodings)
        np.save("model_ids.npy", ids)

        return render_template('entrainement.html', success=True)

    except Exception as e:
        print("❌ Erreur entraînement :", e)
        return render_template('entrainement.html', error=True)
        

# Page de detection faciale
@app.route('/detection_faciale')
def detection_faciale():
    if 'username' not in session:
        return redirect(url_for('login_page'))
    return render_template('detection_faciale.html')

@app.route('/scanner', methods=['POST'])
def scanner():
    try:
        import face_recognition

        #recuperation des model enregistrer dans entrainement
        encodings = np.load("model_encodings.npy", allow_pickle=True)
        ids = np.load("model_ids.npy", allow_pickle=True)

        video = cv2.VideoCapture(0)
        ret, frame = video.read()
        video.release()

        if not ret:
            return render_template('detection_faciale.html', error=True)

        # Recuperation des caracteristiques du visage de la personne a detecter
        face_locations = face_recognition.face_locations(frame)
        face_encs = face_recognition.face_encodings(frame, face_locations)

        conn = get_db_connection()
        cursor = conn.cursor()

        for face_enc in face_encs:

            #calcule de distance entre les valeurs encoder lors de l'entrainement et ceux encoder pendant la detection
            distances = face_recognition.face_distance(encodings, face_enc)

            #recuperation des distance minimal apres le calcul
            min_distance = min(distances)
            index = np.argmin(distances)

            if min_distance < 0.45:         #condition de validation de la reconnaissance
                cne = ids[index]

                # Recuperation des infos du visage detecter pour l'affichage dans 'detection_faciale.html'
                cursor.execute("SELECT NOM FROM ETUDIANT WHERE CNE_MASSAR=%s", (cne,))
                etudiant = cursor.fetchone()

                # Si le visage est introuvable on retourne dans detection_faciale.html pour afficher le message de 
                if not etudiant:
                    return render_template('detection_faciale.html', etudiant_inconnu=True)

                nom = etudiant[0]
                
                # Si l'etudiant a deja été scanner pour la journée on affiche le message 
                cursor.execute("""SELECT * FROM PRESENCE 
                                  WHERE CNE_MASSAR=%s AND DATE_PRESENCE=%s""",
                               (cne, date.today()))
                deja_present = cursor.fetchone()

                # si l'etudiant n'est pas encore été reconnu pour la journée
                if not deja_present:
                    # On recupere les info pour valider la reconnaissance de l'étudiant
                    cursor.execute("""INSERT INTO PRESENCE (CNE_MASSAR, DATE_PRESENCE, HEURE, STATUT)
                                      VALUES (%s, %s, %s, %s)""",
                                   (cne, date.today(), datetime.now().time(), "Présent"))
                    conn.commit()
                    cursor.close()
                    conn.close()
                    return render_template('detection_faciale.html',
                                           presence_enregistree=True,
                                           nom=nom,
                                           cne=cne)
                #si deja reconnu afficher message de presence deja enregistrer
                else:
                    cursor.close()
                    conn.close()
                    return render_template('detection_faciale.html',
                                           presence_deja=True,
                                           nom=nom,
                                           cne=cne)
                
                # Si visage pas reconnu message etudiant_inconnu s'affiche
            else:
                return render_template('detection_faciale.html', etudiant_inconnu=True)

        cursor.close()
        conn.close()
        return render_template('detection_faciale.html', error=True)

    except Exception as e:
        print("❌ Erreur scanner :", e)
        return render_template('detection_faciale.html',
                       presence_enregistree=True,
                       nom=nom,
                       cne=cne)



@app.route('/generer_rapport', methods=['POST', 'GET'])
def generer_rapport():
    conn = get_db_connection()
    cursor = conn.cursor()

    # Requête : toutes les présences, triées par date et nom
    cursor.execute("""
        SELECT e.CNE_MASSAR, e.nom, e.prenom, e.sex, e.options, d.nom_departement, d.chef_departement,
               e.annee_academique, e.semestre, e.mail, e.telephone, p.DATE_PRESENCE, p.HEURE, p.STATUT
        FROM etudiant e
        JOIN departement d ON e.CNE_MASSAR = d.CNE_MASSAR
        JOIN presence p ON e.CNE_MASSAR = p.CNE_MASSAR
        ORDER BY p.DATE_PRESENCE DESC, e.nom ASC
    """)
    presences = cursor.fetchall()

    cursor.close()
    conn.close()

    return render_template('rapport.html', presences=presences)



@app.route('/imprimer_rapport', methods=['POST'])
def imprimer_rapport():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
    SELECT e.CNE_MASSAR, e.nom, e.prenom, e.sexe, e.option, d.nom_departement, d.chef_departement,
           e.annee, e.semestre, e.mail, e.tel, p.DATE_PRESENCE, p.HEURE, p.STATUT
    FROM etudiant e
    JOIN departement d ON e.CNE_MASSAR = d.CNE_MASSAR
    JOIN presence p ON e.CNE_MASSAR = p.CNE_MASSAR
    ORDER BY p.DATE_PRESENCE DESC, e.nom ASC
    """)

    presences = cursor.fetchall()
    cursor.close()
    conn.close()

    doc = Document()
    doc.add_heading('Rapport de Présence', 0)

    table = doc.add_table(rows=1, cols=14)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'CNE'
    hdr_cells[1].text = 'Nom'
    hdr_cells[2].text = 'Prénom'

    for p in presences:
        row_cells = table.add_row().cells
        for i in range(len(p)):
            row_cells[i].text = str(p[i])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name='rapport_presence.docx')


@app.route('/interface_impression', methods=['POST'])
def interface_impression():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT e.CNE_MASSAR, e.nom, e.prenom, e.sex, e.options, d.nom_departement, d.chef_departement,
               e.annee_academique, e.semestre, e.mail, e.telephone, p.DATE_PRESENCE, p.HEURE, p.STATUT
        FROM etudiant e
        JOIN departement d ON e.CNE_MASSAR = d.CNE_MASSAR
        JOIN presence p ON e.CNE_MASSAR = p.CNE_MASSAR
        ORDER BY p.DATE_PRESENCE DESC, e.nom ASC
    """)
    presences = cursor.fetchall()
    cursor.close()
    conn.close()

    return render_template('impression.html', presences=presences, now=datetime.now())


# Déconnexion
@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login_page'))

if __name__ == '__main__':
    app.run(debug=True)
